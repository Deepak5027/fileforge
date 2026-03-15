"""
Document conversion using pure Python libraries only.
PDF reading: pypdf
DOCX reading/writing: python-docx
HTML/MD: html2text + markdown
No LibreOffice, no Pandoc needed.
"""
from pathlib import Path
import shutil


def convert_document(src_path: str, out_path: str, src_fmt: str, tgt_fmt: str):
    src = src_fmt.lower()
    tgt = tgt_fmt.lower()

    # PDF → TXT
    if src == "pdf" and tgt == "txt":
        _pdf_to_txt(src_path, out_path)

    # DOCX → TXT
    elif src == "docx" and tgt == "txt":
        _docx_to_txt(src_path, out_path)

    # DOCX → HTML
    elif src == "docx" and tgt == "html":
        _docx_to_html(src_path, out_path)

    # TXT → MD (passthrough)
    elif src == "txt" and tgt == "md":
        shutil.copy2(src_path, out_path)

    # MD → TXT
    elif src == "md" and tgt == "txt":
        _md_to_txt(src_path, out_path)

    # MD → HTML
    elif src == "md" and tgt == "html":
        _md_to_html(src_path, out_path)

    # HTML → TXT
    elif src in ("html", "htm") and tgt == "txt":
        _html_to_txt(src_path, out_path)

    # TXT → HTML
    elif src == "txt" and tgt == "html":
        _txt_to_html(src_path, out_path)

    else:
        raise ValueError(
            f"Conversion {src} → {tgt} requires LibreOffice which is not available "
            f"on the free tier. Supported free conversions: "
            f"PDF→TXT, DOCX→TXT, DOCX→HTML, TXT↔MD, MD→HTML, HTML→TXT"
        )


def _pdf_to_txt(src_path: str, out_path: str):
    from pypdf import PdfReader
    reader = PdfReader(src_path)
    text = "\n\n".join(
        page.extract_text() or "" for page in reader.pages
    )
    Path(out_path).write_text(text, encoding="utf-8")


def _docx_to_txt(src_path: str, out_path: str):
    from docx import Document
    doc = Document(src_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    Path(out_path).write_text(text, encoding="utf-8")


def _docx_to_html(src_path: str, out_path: str):
    from docx import Document
    doc = Document(src_path)
    lines = ["<html><body>"]
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            level = p.style.name[-1] if p.style.name[-1].isdigit() else "2"
            lines.append(f"<h{level}>{p.text}</h{level}>")
        else:
            lines.append(f"<p>{p.text}</p>")
    lines.append("</body></html>")
    Path(out_path).write_text("\n".join(lines), encoding="utf-8")


def _md_to_txt(src_path: str, out_path: str):
    import html2text
    import markdown
    md_content = Path(src_path).read_text(encoding="utf-8")
    html = markdown.markdown(md_content)
    h = html2text.HTML2Text()
    h.ignore_links = False
    Path(out_path).write_text(h.handle(html), encoding="utf-8")


def _md_to_html(src_path: str, out_path: str):
    import markdown
    md_content = Path(src_path).read_text(encoding="utf-8")
    html = f"<html><body>{markdown.markdown(md_content)}</body></html>"
    Path(out_path).write_text(html, encoding="utf-8")


def _html_to_txt(src_path: str, out_path: str):
    import html2text
    content = Path(src_path).read_text(encoding="utf-8", errors="replace")
    h = html2text.HTML2Text()
    h.ignore_links = False
    Path(out_path).write_text(h.handle(content), encoding="utf-8")


def _txt_to_html(src_path: str, out_path: str):
    content = Path(src_path).read_text(encoding="utf-8", errors="replace")
    lines = content.split("\n")
    html = "<html><body>" + "".join(f"<p>{l}</p>" for l in lines) + "</body></html>"
    Path(out_path).write_text(html, encoding="utf-8")
